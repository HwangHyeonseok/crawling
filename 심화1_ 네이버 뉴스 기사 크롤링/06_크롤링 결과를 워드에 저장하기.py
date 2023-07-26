# -*- coding: euc-kr -*-
# ũ�Ѹ��� ��� ����, �ּ�, ������ ���� ���Ϸ� �����.
# ��� ������ ���忡�� ���� �κ�, �ּҿ� ���� �κ��� ���� ���Ͽ��� ���� �κ����� ������ �Ѵ�.

import requests
from bs4 import BeautifulSoup
import pyautogui
import time
from docx import Document

document = Document() # ���� ���� (��ü) ����

search = pyautogui.prompt("� ���� �˻��Ͻðھ��?")
lastpage = int(pyautogui.prompt("�� ���������� ũ�Ѹ��ұ��?"))
page_num = 1

for now_page in range(1, lastpage * 10, 10): # 5������ ������, 1������ 1 2������ 11 .. 5������ 51
    print(f"{page_num} ������ ũ�Ѹ� ���Դϴ�.")
    response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_pge&query={search}&sort=0&photo=0&field=0&pd=0&ds=&de=&cluster_rank=54&mynews=0&office_type=0&office_section_code=0&news_office_checked=&nso=so:r,p:all,a:all&start={now_page}")
    html = response.text
    soup = BeautifulSoup(html, "html.parser")
    articles = soup.select(".info_group")

    for article in articles:
        # '���̹�����' �� �ִ� ��縸 �����Ѵ�. (<a> �����۸�ũ�� 2�� �̻��� ��쿡 �ش�)
        links = article.select("a.info")
        if len(links) >=2 :
            url = links[1].attrs['href']
            response = requests.get(url, headers={'User-agent':'Mozila/5.0'})
            html = response.text
            soup = BeautifulSoup(html, "html.parser")
            
            # ������ ����� ���
            if "sports" in response.url: # response.url�� ����ϴ� ��� : html �ڵ� �ȿ� �ִ� url ��ũ�� ���� �� �� �������� ��ũ�� �ٸ� ���
                article_title = soup.select_one("h4.title")
                article_body = soup.select_one("#newsEndContents")
                # ���� ���� ���ʿ��� ���� ���� p�±׿� div�±��� ������ ����� �ʿ䰡 ����. ��������.
                p_tags = article_body.select("p") # �������� p �±��� �͵��� ����
                for p_tag in p_tags:
                    p_tag.decompose()

                div_tags = article_body.select("div") # �������� div �±��� �͵��� ����
                for div_tag in div_tags:
                    div_tag.decompose()
            
            # ���� ����� ���
            elif "entertain" in response.url:
                article_title = soup.select_one("h2.end_tit")
                article_body = soup.select_one("#articeBody")
                print(article_title)
                print(article_body)
            
            # �Ϲ� ���� ����� ���
            else:
                article_title = soup.select_one("#title_area")
                article_body = soup.select_one("#dic_area")

            # ��¹�
            print("==================================================== ���� ===========================================================")
            print(article_title.text.strip())
            document.add_heading(article_title.text.strip(), level=0) # ���� ������ �ۼ�
            print("==================================================== �ּ� ===========================================================")
            print(url.strip())
            document.add_paragraph(url.strip()) # ���� ������ �ۼ�
            print("==================================================== ���� ===========================================================")
            print(article_body.text.strip()) #strip �Լ��� �� ���� ������ �����Ѵ�.
            document.add_paragraph(article_body.text.strip()) # ���� ������ �ۼ�
            
            time.sleep(0.3)
    page_num = page_num+1

# ���� ���� �����ϱ�
document.save(r"C:\pratice_crolling\��ȭ1_\article.docx") #�ۼ��� ���� ���� ����