# -*- coding: euc-kr -*-
from docx import Document

# 1. 워드 생성하기
document = Document()

# 2. 워드 데이터 추가하기 # add_heading은 제목, add_paragraph는 본문 내용
document.add_heading('기사 제목', level=0)
document.add_paragraph('기사 링크')
document.add_paragraph('기사 본문')

# 3. 워드 저장하기
document.save(r"C:\pratice_crolling\심화1_\test.docx")