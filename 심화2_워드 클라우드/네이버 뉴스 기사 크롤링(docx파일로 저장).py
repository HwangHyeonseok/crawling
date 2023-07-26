# -*- coding: euc-kr -*-
# -------------------------------------------------- �䱸 ���� ----------------------------------------------------------------
# ���̹� ���� ��縦 ũ�Ѹ� ���ּ���. 
# �䱸 ������ ������ �����ϴ�.
# 1. �˻���� ���̹� ���� ��� �� ���������� ũ�Ѹ� �ϴ��� �Է��� �ް� �ش� ���̹� ���� ������ ũ�Ѹ��Ͽ� ����(.docx Ȯ����)�� �������ּ���.
# 2. ũ�Ѹ��� ������ "���� ������ ���������� ����/���� �Ǿ����ϴ�." ��� �˾�â�� �ߵ��� ���ּ���.
# ------------------------------------------------------------------------------------------------------------------

# ���� ���뿡�� �ֿ��ϰ� ����� Ű���尡 �������� �м��غ���!

import requests
from bs4 import BeautifulSoup
import pyautogui
import time
from openpyxl import Workbook
from openpyxl.styles import Alignment
import pyperclip 
from docx import Document

search = pyautogui.prompt("� ���� �˻��Ͻðھ��?")
lastpage = int(pyautogui.prompt("�� ���������� ũ�Ѹ��ұ��?"))
page_num = 1 # �ֿܼ� ������ ���� ����ϱ� ���� ����
shell_width = 65 # ������ ù ���� 'A' ���� ����
shell_row = 2 # ������ ù ���� 1 ���� �����ε�, �̹� �������� 1�� ä�����Ƿ� 2���� ����

# ��ü ���� ���� ��Ƶδ� ����
total_content = "" 

# ��� ����
article_num = 0

# ���� ���� ����
word = Document()

# ���� ���� ����
excel = Workbook()

# ���� ��ũ��Ʈ ����
ws = excel.create_sheet(f"{search} ��� ����")

# ���� �� �ʺ� ����
ws.column_dimensions['A'].width = 60
ws.column_dimensions['B'].width = 60
ws.column_dimensions['C'].width = 120

# ���� ����κп� ���� �߰�
ws['A1'] = "��� ��ũ"
ws['B1'] = "��� ����"
ws['C1'] = "��� ����"


for now_page in range(1, lastpage * 10, 10): # 5������ ������, 1������ 1 2������ 11 .. 5������ 51
    print(f"{page_num} ������ ũ�Ѹ� ���Դϴ�.")
    response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_pge&query={search}&sort=0&photo=0&field=0&pd=0&ds=&de=&cluster_rank=54&mynews=0&office_type=0&office_section_code=0&news_office_checked=&nso=so:r,p:all,a:all&start={now_page}")
    html = response.text
    soup_naver_search_page = BeautifulSoup(html, "html.parser") # 1��° soup : ���� �˻� ������������ ����
    articles = soup_naver_search_page.select(".info_group")

    for article in articles:
        # '���̹�����' �� �ִ� ��縸 �����Ѵ�. (<a> �����۸�ũ�� 2�� �̻��� ��쿡 �ش�)
        links = article.select("a.info")
        if len(links) >=2 :
            url = links[1].attrs['href']
            response = requests.get(url, headers={'User-agent':'Mozila/5.0'})
            html = response.text
            soup_naver_article_page = BeautifulSoup(html, "html.parser") # 2��° soup : ���̹���縦 ���� ���� �� �������� ����
            
            # ������ ����� ���
            if "sports" in response.url: # response.url�� ����ϴ� ��� : html �ڵ� �ȿ� �ִ� url ��ũ�� ���� �� �� �������� ��ũ�� �ٸ� ���
                article_body = soup_naver_article_page.select_one("#newsEndContents")
                # ���� ���� ���ʿ��� ���� ���� p�±׿� div�±��� ������ ����� �ʿ䰡 ����. ��������.
                p_tags = article_body.select("p") # �������� p �±��� �͵��� ����
                for p_tag in p_tags:
                    p_tag.decompose()

                div_tags = article_body.select("div") # �������� div �±��� �͵��� ����
                for div_tag in div_tags:
                    div_tag.decompose()
            
            # ���� ����� ���
            elif "entertain" in response.url:
                article_body = soup_naver_article_page.select_one("#articeBody")
                print(article_body)
            
            # �Ϲ� ���� ����� ���
            else:
                article_body = soup_naver_article_page.select_one("#dic_area")

            # ��¹�
            print("==================================================== ���� ===========================================================")
            print(article_body.text.strip()) #strip �Լ��� �� ���� ������ �����Ѵ�.
            ws[f"C{shell_row}"] = article_body.text.strip() # ���� ��Ʈ�� ������ �߰�
            total_content += article_body.text.strip() # ���� ������ ��� total_content�� ��� ���� Ŭ����� �м��ϱ� ���� ���
            article_num += 1 # ũ�Ѹ��� ��� ���� ����
            
            # ������ �ڵ� �ٹٲ� ��� ����
            ws[f'C{shell_row}'].alignment = Alignment(wrap_text=True)

            shell_row = shell_row +1 # �� ������ ������ �� ���� �ٷ� �̵��ؼ� �� ����Ѵ�. ex A2 B2 C2 ��� �� �����Ƿ� �״��� �� A3 B3 C3�� ����.
            time.sleep(0.3)
    
    #������ ������ ���� Ȯ���ϱ�
    if(soup_naver_search_page.select_one("a.btn_next").attrs["aria-disabled"] == "true"):
        print("������ �������Դϴ�.")
        break
    page_num = page_num+1

print(f"{article_num} �� ��� ũ�Ѹ� �Ϸ��߽��ϴ�.")
pyperclip.copy(total_content) # ��ü ���� ������ Ŭ������ �ȿ��� �ٿ��־��ش�. (���� ���簡 �Ǿ����Ƿ� ���� Ctrl+V�� �ٿ��ֱ⸸ ���ָ� �ȴ�.)
pyautogui.alert("Ŭ�����忡 ���� ����Ǿ����ϴ�.")

# 2. ���� ������ �߰��ϱ� # add_heading�� ����, add_paragraph�� ���� ����
word.add_paragraph(total_content)
# 3. ���� �����ϱ�
word.save(rf"C:\pratice_crolling\��ȭ2_���� Ŭ����\{search} ��� ũ�Ѹ� ���.docx")
pyautogui.alert("���� ������ ���������� ����/���� �Ǿ����ϴ�.")

excel.save(fr"C:\pratice_crolling\��ȭ1_ ���̹� ���� ��� ũ�Ѹ�\{search} ��� ũ�Ѹ� ���.xlsx")